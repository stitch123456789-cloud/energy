import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# --- 1. 格式與替換核心工具 ---
def set_table_border(table):
    tbl = table._tbl
    ptr = tbl.find(qn('w:tblPr'))
    if ptr is not None:
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            edge = OxmlElement(f'w:{b}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '4') 
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), '000000')
            borders.append(edge)
        ptr.append(borders)

def fix_cell_font(cell, size=12, is_bold=False):
    """鎖定：標楷體、黑色、12號、置中"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = 1 
        if not paragraph.runs: paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(size)
            run.font.bold = is_bold
            run.font.color.rgb = RGBColor(0, 0, 0)

def safe_replace(doc, data_map):
    """處理一般段落與表格內的碎裂標籤替換"""
    # 處理段落
    for p in doc.paragraphs:
        full_text = "".join(run.text for run in p.runs)
        for key, val in data_map.items():
            if key in full_text:
                full_text = full_text.replace(key, str(val))
                p.clear()
                new_run = p.add_run(full_text)
                new_run.font.name = '標楷體'
                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                new_run.font.size = Pt(12)
                new_run.font.color.rgb = RGBColor(0, 0, 0)
    # 處理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = "".join(run.text for run in p.runs)
                    for key, val in data_map.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(val))
                            p.text = full_text
                            for run in p.runs:
                                run.font.name = '標楷體'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                run.font.color.rgb = RGBColor(0, 0, 0)

# --- 2. Streamlit 介面設計 ---
st.title("🌀 P5. 冷卻水塔風車變頻專業分析")

col_a, col_b, col_c = st.columns(3)
with col_a:
    unit_name = st.text_input("單位名稱", value="貴單位")
    ch_info = st.text_input("主機編號", value="CH-1")
with col_b:
    motor_hp = st.number_input("基準馬力 (HP)", value=50.0)
    elec_val = st.number_input("平均電費 (元/度)", value=3.5, step=0.1)
with col_c:
    rt_info = st.text_input("水塔總容量", value="1500RT")
    setup_note = st.text_input("運轉說明", value="僅開啟 2 台")

st.subheader("⚙️ 改善後：各台設備運轉參數調整")
st.info("這裡可以針對每一台風扇設定不同的季節負載與時數（用於產出表二）")

after_rows = []
for t in st.session_state.towers:
    for f in range(1, t['fans'] + 1):
        after_rows.append({
            "設備編號": f"{t['name']}-F{f}",
            "春秋負載(%)": 60, "春秋時數(hr)": 2190,
            "夏季負載(%)": 50, "夏季時數(hr)": 1095,
            "冬季負載(%)": 0,  "冬季時數(hr)": 0
        })

# 讓使用者在網頁上直接改
edit_after_df = st.data_editor(pd.DataFrame(after_rows), use_container_width=True, key="after_editor")
current_op_df = st.data_editor(st.session_state.p5_op_data, use_container_width=True)

if "towers" not in st.session_state:
    st.session_state.towers = [{"name": "CT-1", "rt": 300, "hp": 15.0, "fans": 3}]

with st.sidebar:
    st.header("⚙️ 設備管理")
    if st.button("➕ 新增一組水塔"):
        idx = len(st.session_state.towers) + 1
        st.session_state.towers.append({"name": f"CT-{idx}", "rt": 300, "hp": 15.0, "fans": 1})
        st.rerun()
    if st.button("❌ 刪除最後一組"):
        if len(st.session_state.towers) > 1:
            st.session_state.towers.pop(); st.rerun()

for i, t in enumerate(st.session_state.towers):
    with st.expander(f"配置：{t['name']}", expanded=True):
        tc1, tc2, tc3, tc4 = st.columns(4)
        t['name'] = tc1.text_input("編號", value=t['name'], key=f"n_{i}")
        t['rt'] = tc2.number_input("噸數(RT)", value=t['rt'], key=f"r_{i}")
        t['hp'] = tc3.number_input("馬力(HP)", value=t['hp'], key=f"h_{i}")
        t['fans'] = tc4.number_input("台數", min_value=1, max_value=5, value=t['fans'], key=f"f_{i}")

# --- 3. 生成按鈕與核心邏輯 ---
if st.button("🚀 生成專業效益報告", use_container_width=True):
    try:
        # 1. 核心數據計算
        total_hp = sum(t['hp'] * t['fans'] for t in st.session_state.towers)
        calc_invest = total_hp * 1.3 # 1.3萬/HP
        
        total_old_kwh = 0
        total_new_kwh = 0
        total_kw = total_hp * 0.746
        
        for _, row in current_op_df.iterrows():
            h = float(row["時數(hr)"])
            l = float(row["負載率(%)"]) / 100
            o_kwh = total_kw * h
            n_kwh = total_kw * (l**3) * 1.06 * h
            total_old_kwh += o_kwh
            total_new_kwh += n_kwh
        
        save_kwh = total_old_kwh - total_new_kwh
        save_money = save_kwh * elec_val / 10000
        payback = calc_invest / save_money if save_money > 0 else 0

        doc = Document("template_p5.docx")
        
        # 2. 文字標籤地圖
        data_map = {
            "{{UN}}": unit_name,
            "{{COUNT}}": str(len(st.session_state.towers)),
            "{{CH_INFO}}": ch_info,
            "{{RT_INFO}}": rt_info,
            "{{MT}}": f"{motor_hp}hp",
            "{{ON}}": setup_note,
            "{{OLD_KWH}}": f"{total_old_kwh:,.0f}",
            "{{SAVE_KWH}}": f"{save_kwh:,.0f}",
            "{{MOTOR_SPEC}}": f"{motor_hp}HP x {int(total_hp/motor_hp)}台",
            "{{SAVE_MONEY}}": f"{save_money:.2f}",
            "{{INVEST}}": f"{calc_invest:.1f}",
            "{{PAYBACK}}": f"{payback:.1f}",
            "{{SUPPRESS_KW}}": "13"
        }
        
        safe_replace(doc, data_map)

        for p in doc.paragraphs:
            if "[[OLD_TABLE]]" in p.text: p.text = ""

        # 3. 生成橫向合併表格
        doc.add_page_break()
        doc.add_paragraph("【表一、現況耗電明細分析表 (橫向擴展)】")
        
        fan_list = []
        for t in st.session_state.towers:
            for _ in range(t['fans']):
                fan_list.append({"hp": t['hp'], "rt": t['rt'], "name": t['name']})

        num_cols = 1 + len(fan_list) + 1
        table = doc.add_table(rows=7, cols=num_cols)
        set_table_border(table)

        labels = ["編號", "水塔散熱噸數(RT)", "額定馬力(hp)", "實際耗功(kW)", "全年使用時數(hr)", "負載率(%)", "全年耗電(kWh)"]
        for r, txt in enumerate(labels):
            fix_cell_font(table.cell(r, 0), is_bold=True)
            table.cell(r, 0).text = txt

        col_ptr = 1
        s_kw, s_kwh = 0, 0
        for t in st.session_state.towers:
            f_count = t['fans']
            c_n = table.cell(0, col_ptr).merge(table.cell(0, col_ptr + f_count - 1))
            c_n.text = t['name']; fix_cell_font(c_n, is_bold=True)
            
            c_r = table.cell(1, col_ptr).merge(table.cell(1, col_ptr + f_count - 1))
            c_r.text = f"{t['rt']}RT"; fix_cell_font(c_r)

            for i in range(f_count):
                kw_f = t['hp'] * 0.746
                kwh_f = kw_f * 4380
                table.cell(2, col_ptr + i).text = f"{t['hp']:.1f}"
                table.cell(3, col_ptr + i).text = f"{kw_f:.1f}"
                table.cell(4, col_ptr + i).text = "4,380"
                table.cell(5, col_ptr + i).text = "100%"
                table.cell(6, col_ptr + i).text = f"{kwh_f:,.0f}"
                s_kw += kw_f; s_kwh += kwh_f
                for r in range(2, 7): fix_cell_font(table.cell(r, col_ptr + i))
            col_ptr += f_count
# --- [新增] 生成改善後節能效益明細表 ---
        doc.add_paragraph("\n【表二、改善後變頻節能效益分析表】")
        
        # 欄數與改善前一致 (標題欄 + 風扇台數 + 合計欄)
        table_after = doc.add_table(rows=20, cols=num_cols)
        set_table_border(table_after)

        # 1. 填寫左側標題標籤
        after_labels = [
            "編號", "春秋季負載率(%)", "夏季負載率(%)", "冬季負載率(%)",
            "春秋季使用時數(hr)", "夏季使用時數(hr)", "冬季使用時數(hr)",
            "春秋季耗功(kW)", "夏季耗功(kW)", "冬季耗功(kW)",
            "春秋季耗電(kWh)", "夏季耗電(kWh)", "冬季耗電(kWh)",
            "改善後總耗電(kWh)", "節省耗電(kWh)", "抑低需量(kW)",
            "節約金額(萬元/年)", "節能比(%)", "註1：變頻器損失", "註2：平均電費(元/kWh)"
        ]
        for r, txt in enumerate(after_labels):
            table_after.cell(r, 0).text = txt
            fix_cell_font(table_after.cell(r, 0), is_bold=True)

        # 2. 準備季節參數 (從介面 current_op_df 抓取)
        # 假設 0:春秋, 1:夏, 2:冬
        s_spring = current_op_df.iloc[0]
        s_summer = current_op_df.iloc[1]
        s_winter = current_op_df.iloc[2]

        # 3. 填寫數據
        col_ptr = 1
        total_after_kwh = 0
        total_save_kwh = 0
        total_suppress_kw = 0
        
        for t in st.session_state.towers:
            f_count = int(t['fans'])
            # 合併編號格
            c_n_after = table_after.cell(0, col_ptr).merge(table_after.cell(0, col_ptr + f_count - 1))
            c_n_after.text = t['name']
            fix_cell_font(c_n_after, is_bold=True)

            kw_base = t['hp'] * 0.746
            for i in range(f_count):
                cur_c = col_ptr + i
                
                # 計算各季變頻功耗 (Load^3 * 1.06)
                spring_kw = kw_base * ((s_spring['負載率(%)']/100)**3) * 1.06
                summer_kw = kw_base * ((s_summer['負載率(%)']/100)**3) * 1.06
                winter_kw = kw_base * ((s_winter['負載率(%)']/100)**3) * 1.06
                
                # 各季耗電
                spring_kwh = spring_kw * s_spring['時數(hr)']
                summer_kwh = summer_kw * s_summer['時數(hr)']
                winter_kwh = winter_kw * s_winter['時數(hr)']
                
                # 該台風扇節電量 (改善前固定 Load=100% 耗電 - 改善後各季總和)
                old_kwh_fan = kw_base * (s_spring['時數(hr)'] + s_summer['時數(hr)'] + s_winter['時數(hr)'])
                after_kwh_fan = spring_kwh + summer_kwh + winter_kwh
                save_kwh_fan = old_kwh_fan - after_kwh_fan
                
                # 填入格子
                table_after.cell(1, cur_c).text = f"{s_spring['負載率(%)']}%"
                table_after.cell(2, cur_c).text = f"{s_summer['負載率(%)']}%"
                table_after.cell(3, cur_c).text = f"{s_winter['負載率(%)']}%"
                table_after.cell(4, cur_c).text = f"{s_spring['時數(hr)']:,.0f}"
                table_after.cell(5, cur_c).text = f"{s_summer['時數(hr)']:,.0f}"
                table_after.cell(6, cur_c).text = f"{s_winter['時數(hr)']:,.0f}"
                table_after.cell(7, cur_c).text = f"{spring_kw:.1f}"
                table_after.cell(8, cur_c).text = f"{summer_kw:.1f}"
                table_after.cell(9, cur_c).text = f"{winter_kw:.1f}"
                table_after.cell(10, cur_c).text = f"{spring_kwh:,.0f}"
                table_after.cell(11, cur_c).text = f"{summer_kwh:,.0f}"
                table_after.cell(12, cur_c).text = f"{winter_kwh:,.0f}"
                table_after.cell(13, cur_c).text = f"{after_kwh_fan:,.0f}"
                table_after.cell(14, cur_c).text = f"{save_kwh_fan:,.0f}"
                table_after.cell(15, cur_c).text = f"{(kw_base * 0.15):.1f}" # 抑低需量估計
                
                # 累加總計
                total_after_kwh += after_kwh_fan
                total_save_kwh += save_kwh_fan
                total_suppress_kw += (kw_base * 0.15)
                
                for r in range(1, 16): fix_cell_font(table_after.cell(r, cur_c))
            col_ptr += f_count

        # 4. 填寫合計與下方合併欄位
        table_after.cell(0, num_cols-1).text = "合計"
        # ... (中間各季加總可視需要補上)
        table_after.cell(13, num_cols-1).text = f"{total_after_kwh:,.0f}"
        table_after.cell(14, num_cols-1).text = f"{total_save_kwh:,.0f}"
        table_after.cell(15, num_cols-1).text = f"{total_suppress_kw:.1f}"
        for r in [0, 13, 14, 15]: fix_cell_font(table_after.cell(r, num_cols-1), is_bold=True)

        # 合併底部的全寬欄位
        for row_idx in range(16, 20):
            merged_cell = table_after.cell(row_idx, 1).merge(table_after.cell(row_idx, num_cols-1))
            if row_idx == 16: merged_cell.text = f"{(total_save_kwh * elec_val / 10000):.1f}"
            elif row_idx == 17: merged_cell.text = f"{(total_save_kwh / (total_after_kwh + total_save_kwh) * 100):.1f}%"
            elif row_idx == 18: merged_cell.text = "6%"
            elif row_idx == 19: merged_cell.text = f"{elec_val:.2f}"
            fix_cell_font(merged_cell, is_bold=True)
        # 合計欄
        table.cell(0, num_cols-1).text = "合計"
        table.cell(3, num_cols-1).text = f"{s_kw:.1f}"
        table.cell(6, num_cols-1).text = f"{s_kwh:,.0f}"
        for r in [0, 3, 6]: fix_cell_font(table.cell(r, num_cols-1), is_bold=True)

        buf = io.BytesIO()
        doc.save(buf)
        st.success("✅ 報告已生成！")
        st.download_button("📥 下載完整報告", buf.getvalue(), "風車分析報告.docx")

    except Exception as e:
        st.error(f"❌ 錯誤: {e}")
