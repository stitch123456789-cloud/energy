import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- 1. 字體與格式函數 ---
def set_font_kai_11(run, color_red=False):
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(11)
    run.font.bold = False
    if color_red:
        run.font.color.rgb = RGBColor(255, 0, 0)

def set_font_kai_bold_14(run):
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(14)
    run.font.bold = True

# --- 2. 數據抓取：照明系統 (全域加總) ---
def fetch_and_aggregate_lighting(file):
    try:
        xl = pd.ExcelFile(file)
        target_sheets = [s for s in xl.sheet_names if "表九之二" in s]
        if not target_sheets: return None
        
        aggregated_data = {}
        for sheet in target_sheets:
            df = pd.read_excel(file, sheet_name=sheet, header=None)
            for i in range(6, len(df)):
                kind = str(df.iloc[i, 1]).strip() # B
                spec = str(df.iloc[i, 5]).strip() # F
                count_str = str(df.iloc[i, 9]).strip()  # J
                hours_str = str(df.iloc[i, 11]).strip() # L
                
                if kind == "nan" or "註" in kind or "合計" in kind: continue
                if spec == "nan" or spec == "": continue
                if '.' in kind: kind = kind.split('.')[-1].strip()
                
                try:
                    count = int(float(count_str.replace(',', '')))
                    hours = int(float(hours_str.replace(',', '')))
                    key = (kind, spec, hours)
                    aggregated_data[key] = aggregated_data.get(key, 0) + count
                except: continue
        return aggregated_data
    except: return None

# --- 3. 數據抓取：(2) 冰水主機規格 (精準解析) ---
def fetch_chiller_spec(file):
    try:
        xl = pd.ExcelFile(file)
        # 精確過濾：必須包含「九之一」而且包含「(三)」
        target_sheets = [s for s in xl.sheet_names if "空調系統(三)" in s]
        if not target_sheets: return None
        
        all_chillers = []
        for sheet in target_sheets:
            df = pd.read_excel(file, sheet_name=sheet, header=None)
            for i in range(6, len(df)):
                name_raw = str(df.iloc[i, 1]).strip() # B: 設備名稱
                if "主機" not in name_raw: continue
                # --- 清洗：去除名稱前方的「1. 」數字 ---
                name = name_raw.split('.')[-1].strip() if '.' in name_raw else name_raw
                sn = str(df.iloc[i, 2]).strip()      # C: 編號
                form = str(df.iloc[i, 5]).strip()    # F: 型式
                inverter_raw = str(df.iloc[i, 7]).strip() # H: 有無 (變頻)
                volt = str(df.iloc[i, 11]).strip()   # L: 電壓
                power = str(df.iloc[i, 12]).strip()  # M: 功率值
                year = str(df.iloc[i, 13]).strip()   # N: 年份
                cap_val = str(df.iloc[i, 14]).strip() # O: 容量
                cap_unit = str(df.iloc[i, 15]).strip() # P: 單位
                qty = str(df.iloc[i, 21]).strip()    # V: 數量
                
                if cap_val == "nan" or cap_val == "": continue

                # 單位換算 RT
                try:
                    val = float(cap_val.replace(',', ''))
                    if "kW" in cap_unit.upper(): rt_val = round(val / 3.517, 1)
                    elif "KCAL" in cap_unit.upper(): rt_val = round(val / 3024, 1)
                    else: rt_val = val
                except: rt_val = cap_val
                
                type_tag = "變頻" if inverter_raw == "有" else "定頻"
                clean_year = year.replace('民國', '').replace('年', '').strip()
                
                all_chillers.append([name, sn, form, volt, power, clean_year, rt_val, qty, type_tag])
        return all_chillers
    except: return None
def fetch_pump_data(file):
    try:
        xl = pd.ExcelFile(file)
        target_sheets = [s for s in xl.sheet_names if "空調系統(三)" in s]
        if not target_sheets: return None, False
        
        pumps = {"冰水泵": [], "區域水泵": [], "冷卻水泵": []}
        has_secondary = False # 判定是否有區域水泵
        
        for sheet in target_sheets:
            df = pd.read_excel(file, sheet_name=sheet, header=None)
            for i in range(6, len(df)):
                name = str(df.iloc[i, 1]).strip() # B: 設備名稱
                # 判定是否為泵浦
                target_key = None
                if "冰水泵" in name: target_key = "冰水泵"
                elif "區域水泵" in name: 
                    target_key = "區域水泵"
                    has_secondary = True # 偵測到區域水泵，觸發 /二次 顯示
                elif "冷卻水泵" in name: target_key = "冷卻水泵"
                
                if not target_key: continue
                
                sn = str(df.iloc[i, 2]).strip()      # C: 編號
                inv_raw = str(df.iloc[i, 7]).strip() # H: 變頻判定
                flow_val = str(df.iloc[i, 14]).strip() # O: 流量數值
                flow_unit = str(df.iloc[i, 15]).strip()# P: 流量單位
                hp_val = str(df.iloc[i, 18]).strip()   # S: 馬力 (HP)
                qty = str(df.iloc[i, 21]).strip()    # V: 數量
                
                if flow_val == "nan" or hp_val == "nan": continue

                # A. 流量單位換算 (GPM -> LPM)
                try:
                    f_val = float(flow_val.replace(',', ''))
                    lpm = round(f_val * 3.785, 0) if "GPM" in flow_unit.upper() else f_val
                    hp = float(hp_val)
                except: continue

                # B. 揚程推算 (馬達效率 62%)
                # 公式: H = (HP * 746 * 0.62) / ( (LPM/60) * 9.8 )
                try:
                    head = round((hp * 462.52) / (lpm / 60 * 9.8), 1) if lpm > 0 else 0
                except: head = 0
                
                type_tag = "變頻" if inv_raw == "有" else "定頻"
                pumps[target_key].append([sn, int(hp), int(lpm), head, qty, type_tag])
                
        return pumps, has_secondary
    except: return None, False
def fetch_cooling_system_data(file):
    try:
        xl = pd.ExcelFile(file)
        target_sheets = [s for s in xl.sheet_names if "空調系統(三)" in s]
        cooling_data = {"冷卻水泵": [], "冷卻水塔": []}
        
        for sheet in target_sheets:
            df = pd.read_excel(file, sheet_name=sheet, header=None)
            for i in range(6, len(df)):
                name = str(df.iloc[i, 1]).strip() # B: 設備名稱
                sn = str(df.iloc[i, 2]).strip()   # C: 編號
                inv = "變頻" if str(df.iloc[i, 7]).strip() == "有" else "定頻"
                cap_val = str(df.iloc[i, 14]).strip() # O: 流量或容量
                unit = str(df.iloc[i, 15]).strip().upper() # P: 單位
                hp_val = str(df.iloc[i, 18]).strip()  # S: 馬力
                qty = str(df.iloc[i, 21]).strip()     # V: 數量

                if cap_val == "nan" or hp_val == "nan": continue

                # --- 處理冷卻水泵 ---
                if "冷卻水泵" in name:
                    try:
                        f_lpm = float(cap_val.replace(',','')) * 3.785 if "GPM" in unit else float(cap_val.replace(',',''))
                        hp = float(hp_val)
                        # 揚程推算 (效率62%)
                        head = round((hp * 462.52) / (f_lpm / 60 * 9.8), 1) if f_lpm > 0 else 0
                        cooling_data["冷卻水泵"].append([sn, int(hp), int(f_lpm), head, qty, inv])
                    except: continue

                # --- 處理冷卻水塔 ---
                elif "冷卻水塔" in name:
                    # 水塔馬力通常顯示 10*3 這種形式，如果 Excel 裡有 * 就直接帶入字串
                    cooling_data["冷卻水塔"].append([sn, hp_val, cap_val, qty, inv])
        
        return cooling_data
    except: return None
# --- 4. Word 生成函數 ---

def add_lighting_table(doc, lighting_data):
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("2. 照明系統：")
    set_font_kai_bold_14(run_title)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 3).merge(table.cell(1, 3))
    
    headers = [(0,0,"燈具種類"),(0,1,"燈具形式"),(0,3,"運轉時數(小時/年)"),(1,1,"容量規格"),(1,2,"數量")]
    for r, c, txt in headers:
        cell = table.cell(r, c)
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt); set_font_kai_11(run)

    sorted_items = sorted(lighting_data.items(), key=lambda x: x[0][0])
    for (kind, spec, hours), count in sorted_items:
        row_cells = table.add_row().cells
        for idx, val in enumerate([kind, spec, str(count), str(hours)]):
            p = row_cells[idx].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val); set_font_kai_11(run)

def add_ac_mode_table(doc, ac_data):
    p3 = doc.add_paragraph()
    run3 = p3.add_run("3. 空調系統："); set_font_kai_bold_14(run3)
    p1 = doc.add_paragraph()
    p1.paragraph_format.left_indent = Pt(20)
    run1 = p1.add_run("(1) 空調主機開啟模式："); set_font_kai_bold_14(run1)

    table = doc.add_table(rows=4, cols=6); table.style = 'Table Grid'
    headers = ["季節", "主機總容量\n(RT)", "冰機總開啟台數", "負載率\n(%)", "合計容量\n(RT)", "出水溫度設定 (°C)"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i); cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); set_font_kai_11(run)

    for r_idx, row_vals in enumerate(ac_data, start=1):
        for c_idx, val in enumerate(row_vals):
            cell = table.cell(r_idx, c_idx); cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val)); set_font_kai_11(run)

def add_chiller_spec_table(doc, chiller_data):
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Pt(20)
    run2 = p2.add_run("(2) 冰水主機規格："); set_font_kai_bold_14(run2)

    table = doc.add_table(rows=1, cols=9); table.style = 'Table Grid'
    headers = ["設備名稱", "設備編號", "型式", "電壓\n(V)", "功率值\n(kW)", "製造年份", "容量\n(RT)", "現有數量", "備註"]
    for i, h in enumerate(headers):
        p = table.cell(0, i).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); set_font_kai_11(run)

    for row_data in chiller_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            p = row_cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val).replace('.0', ''))
            set_font_kai_11(run) # 此處已修正為預設黑字
def add_pump_section(doc, pump_data, has_secondary):
    # (3) 冰水管路系統
    p3 = doc.add_paragraph(); p3.paragraph_format.left_indent = Pt(20)
    set_font_kai_bold_14(p3.add_run("(3) 冰水管路系統："))
    
    side_txt = "採一/二次側系統設計" if has_sec else "採一次側系統設計"
    ps = doc.add_paragraph(); ps.paragraph_format.left_indent = Pt(40)
    set_font_kai_11(ps.add_run(f"{side_txt}，設備規格說明如下表所示"))

    for p_name in ["冰水泵", "區域水泵", "冷卻水泵"]:
        items = pump_data.get(p_name, [])
        if not items: continue
        
        # 建立表格 (多加一列放標題，共 3 列標頭 + 數據)
        # 初始 2 列：第一列標題，第二列欄位名稱，數據在後續 add_row
        table = doc.add_table(rows=2, cols=6); table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- 第一列：小標題 (冰水泵等) ---
        title_cell = table.cell(0, 0).merge(table.cell(0, 5))
        cp_title = title_cell.paragraphs[0]; cp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_kai_11(cp_title.add_run(p_name))

        # --- 第二列：欄位名稱 ---
        h3 = ["設備編號", "額定馬力\n(HP)", "額定流量\n(LPM)", "揚程\n(m)", "數量\n(台)", "備註"]
        for i, txt in enumerate(h3):
            cp = table.cell(1, i).paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font_kai_11(cp.add_run(txt))

        # --- 填入數據 ---
        for r_vals in items:
            row = table.add_row().cells
            for i, v in enumerate(r_vals):
                cp = row[i].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 將數值轉為字串並去除 .0
                clean_v = str(v).replace('.0', '')
                set_font_kai_11(cp.add_run(clean_v))
        
        doc.add_paragraph() # 每個表格間的間距
        def add_cooling_section(doc, cooling_data):
    # (4) 冷卻水系統
    p4 = doc.add_paragraph(); p4.paragraph_format.left_indent = Pt(20)
    set_font_kai_bold_14(p4.add_run("(4) 冷卻水系統："))

    # --- 冷卻水泵表格 ---
    if cooling_data["冷卻水泵"]:
        table = doc.add_table(rows=2, cols=6); table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 標籤列
        title_cell = table.cell(0, 0).merge(table.cell(0, 5))
        set_font_kai_11(title_cell.paragraphs[0].add_run("冷卻水泵"))
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 欄位名稱
        h_pumps = ["設備編號", "額定馬力\n(HP)", "額定流量\n(LPM)", "揚程\n(m)", "數量\n(台)", "備註"]
        for i, txt in enumerate(h_pumps):
            set_font_kai_11(table.cell(1, i).paragraphs[0].add_run(txt))
            table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 填值
        for r_vals in cooling_data["冷卻水泵"]:
            row = table.add_row().cells
            for i, v in enumerate(r_vals):
                set_font_kai_11(row[i].paragraphs[0].add_run(str(v).replace('.0','')))
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # --- 冷卻水塔表格 ---
    if cooling_data["冷卻水塔"]:
        table = doc.add_table(rows=2, cols=5); table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 標籤列
        title_cell = table.cell(0, 0).merge(table.cell(0, 4))
        set_font_kai_11(title_cell.paragraphs[0].add_run("冷卻水塔"))
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 欄位名稱
        h_tower = ["設備編號", "額定馬力\n(HP)", "額定容量\n(RT)", "數量\n(台)", "備註"]
        for i, txt in enumerate(h_tower):
            set_font_kai_11(table.cell(1, i).paragraphs[0].add_run(txt))
            table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 填值
        for r_vals in cooling_data["冷卻水塔"]:
            row = table.add_row().cells
            for i, v in enumerate(r_vals):
                set_font_kai_11(row[i].paragraphs[0].add_run(str(v).replace('.0','')))
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # (5) 空調附屬設備 (固定文字)
    p5 = doc.add_paragraph(); p5.paragraph_format.left_indent = Pt(20)
    set_font_kai_bold_14(p5.add_run("(5) 空調附屬設備："))
    run5 = p5.add_run("空氣側採用空調箱或小型送風機供應空調至現場使用。")
    set_font_kai_11(run5)
# --- 5. Streamlit 介面 ---
st.subheader("⚙️ 設備系統資料庫")

st.markdown("### ❄️ 3. 空調主機開啟模式設定")
c0, c1, c2, c3, c4 = st.columns([0.8, 1.5, 1.2, 1.2, 1.5])
with c0:
    st.write("**季節**")
    st.markdown("<br>", unsafe_allow_html=True); st.caption("夏季")
    st.markdown("<br>", unsafe_allow_html=True); st.caption("春秋")
    st.markdown("<br>", unsafe_allow_html=True); st.caption("冬季")
with c1:
    st.write("**主機總容量(RT)**")
    rt_s = st.number_input("夏量", 600, key="v_rt_s", label_visibility="collapsed")
    rt_sp = st.number_input("春量", 450, key="v_rt_sp", label_visibility="collapsed")
    rt_w = st.number_input("冬量", 450, key="v_rt_w", label_visibility="collapsed")
with c2:
    st.write("**台數**")
    ct_s = st.number_input("夏台", 1, key="v_ct_s", label_visibility="collapsed")
    ct_sp = st.number_input("春台", 1, key="v_ct_sp", label_visibility="collapsed")
    ct_w = st.number_input("冬台", 1, key="v_ct_w", label_visibility="collapsed")
with c3:
    st.write("**負載率(%)**")
    ld_s = st.number_input("夏負", 70, key="v_ld_s", label_visibility="collapsed")
    ld_sp = st.number_input("春負", 70, key="v_ld_sp", label_visibility="collapsed")
    ld_w = st.number_input("冬負", 60, key="v_ld_w", label_visibility="collapsed")
with c4:
    st.write("**出水溫度(°C)**")
    tp_s = st.number_input("夏溫", 7, key="v_tp_s", label_visibility="collapsed")
    tp_sp = st.number_input("春溫", 7, key="v_tp_sp", label_visibility="collapsed")
    tp_w = st.number_input("冬溫", 7, key="v_tp_w", label_visibility="collapsed")

ac_rows = [
    ["夏季", rt_s, ct_s, f"{ld_s}%", round(rt_s*ld_s/100, 1), tp_s],
    ["春秋", rt_sp, ct_sp, f"{ld_sp}%", round(rt_sp*ld_sp/100, 1), tp_sp],
    ["冬季", rt_w, ct_w, f"{ld_w}%", round(rt_w*ld_w/100, 1), tp_w]
]

st.markdown("---")
up_file = st.file_uploader("請上傳能源查核 Excel", type=["xlsx"])
final_file = up_file if up_file else st.session_state.get('global_excel')

if final_file:
    if st.button("🚀 生成並下載設備系統報告", use_container_width=True):
        doc = Document()
        
        # 1. 處理照明 (確保這段在按鈕內)
        l_data = fetch_and_aggregate_lighting(final_file)
        if l_data: add_lighting_table(doc, l_data)
        
        # 2. 處理空調開啟模式
        add_ac_mode_table(doc, ac_rows)
        
        # 3. 處理冰水主機
        c_data = fetch_chiller_spec(final_file)
        if c_data: add_chiller_spec_table(doc, c_data)
        
        # 4. 處理管路與冷卻水系統 (確保這些都在按鈕內)
        p_data, has_sec = fetch_pump_and_cooling_data(final_file)
        if p_data:
            add_pump_section(doc, p_data, has_sec)    # (3) 冰水管路
            add_cooling_section(doc, p_data)         # (4) 冷卻水系統 & (5) 附屬設備
        
        # 5. 下載 (這段必須在按鈕內的最末尾)
        buf = io.BytesIO()
        doc.save(buf)
        st.download_button("📥 下載 Word 報告", buf.getvalue(), "設備報告.docx", use_container_width=True)
