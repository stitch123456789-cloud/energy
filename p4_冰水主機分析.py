import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- 1. 字體工具函數 ---
def add_run_kai(paragraph, text, size=12, is_bold=False):
    run = paragraph.add_run(str(text)) # 確保 text 是字串防止 Attribute 錯誤
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run

# --- 2. 初始值與 Session State 邏輯 ---
default_base_old = 0.95
default_base_new = 0.50

def reset_all_data():
    keys_to_clear = [
        "old_cfg_data", "new_cfg_data", "old_op_data", "new_op_data",
        "old_cfg_edit", "new_cfg_edit", "old_op_edit", "new_op_edit"
    ]
    for k in keys_to_clear:
        if k in st.session_state:
            del st.session_state[k]
    st.rerun()

if "old_cfg_data" not in st.session_state:
    st.session_state.old_cfg_data = pd.DataFrame([{"編號": "CH-1", "台數": 2, "容量(RT)": 500, "型式": "螺旋式"}])

if "new_cfg_data" not in st.session_state:
    st.session_state.new_cfg_data = st.session_state.old_cfg_data.copy()
    st.session_state.new_cfg_data.at[0, "型式"] = "離心式"

if "old_op_data" not in st.session_state:
    st.session_state.old_op_data = pd.DataFrame({
        "季節": ["春秋", "夏季", "冬季"],
        "RT": [500, 500, 500],
        "台數": [1, 1, 1],
        "時數(hr/y)": [4380, 1095, 1095],
        "負載率(%)": [70, 80, 50],
        "效率(kW/RT)": [round(default_base_old*0.96,3), default_base_old, round(default_base_old*0.94,3)]
    })

if "new_op_data" not in st.session_state:
    st.session_state.new_op_data = st.session_state.old_op_data.copy()
    st.session_state.new_op_data["效率(kW/RT)"] = [round(default_base_new*0.96,3), base_new_eff if 'base_new_eff' in locals() else default_base_new, round(default_base_new*0.94,3)]

# --- 3. Streamlit 介面佈局 ---
st.title("❄️ P4. 冰水主機汰換效益分析")

c1, c2, c3 = st.columns(3)
with c1:
    unit_name = st.text_input("單位名稱", value="貴單位")
    if st.button("♻️ 重置所有表格資料", use_container_width=True):
        reset_all_data()
with c2:
    val_from_app = st.session_state.get('auto_avg_price', 4.48)
    elec_price = st.number_input("平均電費 (元/度)", value=float(val_from_app), step=0.01)
with c3:
    setup_year = st.number_input("原主機設置年份 (民國)", value=104)

st.markdown("---")
st.subheader("🎯 效率與名稱快速設定")
ca, cb, cc = st.columns(3)
with ca:
    suggest_ch_name = st.text_input("建議更換主機名稱", value="CH-1")
with cb:
    base_old_eff = st.number_input("現況夏季效率基準 (kW/RT)", value=default_base_old, step=0.01)
with cc:
    base_new_eff = st.number_input("改善後夏季效率基準 (kW/RT)", value=default_base_new, step=0.01)

# --- 4. 介面表格與連動邏輯 ---
st.markdown("---")
left_col, right_col = st.columns(2)

with left_col:
    st.subheader("🧊 1. 改善前 (現況)")
    old_cfg_raw = st.data_editor(st.session_state.old_cfg_data, num_rows="dynamic", use_container_width=True, key="old_cfg_edit")
    old_op_raw = st.data_editor(st.session_state.old_op_data, use_container_width=True, key="old_op_edit")

    # 過濾空行後的乾淨資料
    old_cfg = old_cfg_raw.dropna(how='all').dropna(subset=['台數', '容量(RT)'])
    old_op = old_op_raw.dropna(how='all').dropna(subset=['RT', '台數', '效率(kW/RT)'])

    if not old_op_raw.equals(st.session_state.old_op_data):
        st.session_state.old_op_data = old_op_raw
        for col in ["RT", "台數", "時數(hr/y)", "負載率(%)"]:
            st.session_state.new_op_data[col] = old_op_raw[col]
        st.rerun()

with right_col:
    st.subheader("✨ 2. 改善後 (預期)")
    new_cfg_raw = st.data_editor(st.session_state.new_cfg_data, num_rows="dynamic", use_container_width=True, key="new_cfg_edit")
    new_op_raw = st.data_editor(st.session_state.new_op_data, use_container_width=True, key="new_op_edit")
    
    # 過濾空行
    new_cfg = new_cfg_raw.dropna(how='all').dropna(subset=['台數', '容量(RT)'])
    new_op = new_op_raw.dropna(how='all').dropna(subset=['RT', '台數', '效率(kW/RT)'])
    
    st.session_state.new_op_data = new_op_raw
    st.session_state.new_cfg_data = new_cfg_raw

st.markdown("---")
invest_amount = st.number_input("請輸入預估投資金額 (萬元)", value=1050)

# --- 5. Word 生成與計算 ---
def build_word_table(doc, op_df_clean):
    table = doc.add_table(rows=1, cols=7); table.style = 'Table Grid'
    hd = ["季節", "製冷量\n(RT)", "台數", "運轉耗電率\n(kW/RT)", "時數\n(時/年)", "負載率", "耗電\n(kWh/年)"]
    for i, h in enumerate(hd):
        cp = table.cell(0,i).paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_kai(cp, h, size=10, is_bold=True)

    total_kwh = 0
    for _, row in op_df_clean.iterrows():
        # 計算前做型態檢查
        rt = float(row["RT"])
        qty = float(row["台數"])
        eff = float(row["效率(kW/RT)"])
        hrs = float(row["時數(hr/y)"])
        load = float(row["負載率(%)"]) / 100
        
        kwh = rt * qty * eff * hrs * load
        total_kwh += kwh
        
        r_cells = table.add_row().cells
        vals = [row["季節"], f"{rt:,.0f}", f"{qty:,.0f}", f"{eff:.3f}", f"{hrs:,.0f}", f"{row['負載率(%)']}%", f"{kwh:,.0f}"]
        for i, v in enumerate(vals):
            cp = r_cells[i].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_kai(cp, v)
    
    row_sum = table.add_row().cells
    row_sum[0].merge(row_sum[5])
    add_run_kai(row_sum[0].paragraphs[0], "總耗電量(kWh/年)", is_bold=True)
    row_sum[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_kai(row_sum[6].paragraphs[0], f"{total_kwh:,.0f}", is_bold=True)
    row_sum[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return total_kwh

doc = Document()
add_run_kai(doc.add_heading('', level=1), "一、現況說明", size=14, is_bold=True)

# 這裡避免空配置導致報錯
if not old_cfg.empty:
    old_desc = "、".join([f"{r['台數']}台{r['容量(RT)']}RT {r['型式']}" for _, r in old_cfg.iterrows()])
else:
    old_desc = "(未設定主機)"

p1 = doc.add_paragraph(); p1.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p1, f"1. {unit_name}空調系統有{old_desc}冰水主機(設置年份{setup_year}年)，推估年度耗電量如下表：")

total_old_kwh = build_word_table(doc, old_op)
p_old_sum = doc.add_paragraph(); p_old_sum.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p_old_sum, f"2.推估耗電量：{total_old_kwh:,.0f} kWh/年。")

doc.add_paragraph()
add_run_kai(doc.add_heading('', level=1), "二、改善方案", size=14, is_bold=True)
p2 = doc.add_paragraph(); p2.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p2, f"1. 建議編列經費汰換為高效率冰水主機，目前新型高效率 1 級能效離心式冰水主機之運轉效率可達 {base_new_eff:.2f} kW/RT，如與以上大樓現況冰水主機運轉效率相比，有節能空間。")
p3 = doc.add_paragraph(); p3.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p3, f"2.參考(附件A-4)『冰水機組製冷能源效率分級基準表』，擬建議貴單位優先將現況低效率之冰水主機{suggest_ch_name}，汰換為符合建議標準的冰水主機，以節省主機運轉耗能。")

doc.add_paragraph()
add_run_kai(doc.add_heading('', level=1), "三、預期效益", size=14, is_bold=True)
p_res_title = doc.add_paragraph(); p_res_title.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p_res_title, "改善後冰水主機耗電量計算如下：")

if not new_cfg.empty:
    new_desc_word = " + ".join([f"{r['容量(RT)']}RT×{r['台數']}" for _, r in new_cfg.iterrows()])
else:
    new_desc_word = "0"
p5 = doc.add_paragraph(); p5.paragraph_format.first_line_indent = Pt(24)
add_run_kai(p5, f"1. 採用高效率離心式冰水主機 {new_desc_word} 台，推估年度耗電量如下表：")

total_new_kwh = build_word_table(doc, new_op)

# 效益計算與表格
save_kwh = total_old_kwh - total_new_kwh
save_rate = (save_kwh / total_old_kwh * 100) if total_old_kwh > 0 else 0
save_money = save_kwh * elec_price / 10000

# 抑制需量安全抓取
try:
    s_old = old_op.iloc[1]
    s_new = new_op.iloc[1]
    suppress_demand = (float(s_old['RT'])*float(s_old['台數'])*float(s_old['效率(kW/RT)'])) - (float(s_new['RT'])*float(s_new['台數'])*float(s_new['效率(kW/RT)']))
except:
    suppress_demand = 0

payback_year = (invest_amount / save_money) if save_money > 0 else 0

doc.add_paragraph()
summary_table = doc.add_table(rows=1, cols=8); summary_table.style = 'Table Grid'
s_hd = ["改善前\n(kWh/年)", "改善後\n(kWh/年)", "節約度數\n(kWh/年)", "節能率\n(%)", "節能電費\n(萬元/年)", "抑制需量\n(kW)", "投資金額\n(萬元)", "回收年限\n(年)"]
for i, h in enumerate(s_hd):
    cp = summary_table.cell(0,i).paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_kai(cp, h, size=9, is_bold=True)

s_row = summary_table.add_row().cells
s_vals = [f"{total_old_kwh:,.0f}", f"{total_new_kwh:,.0f}", f"{save_kwh:,.0f}", f"{save_rate:.1f}", f"{save_money:.1f}", f"{suppress_demand:.1f}", f"{invest_amount:,.0f}", f"{payback_year:.1f}"]
for i, v in enumerate(s_vals):
    cp = s_row[i].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_kai(cp, v)

note_row = summary_table.add_row().cells; note_row[0].merge(note_row[7])
add_run_kai(note_row[0].paragraphs[0], f"註：每度電平均單價 {elec_price:.2f} 元").alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()
add_run_kai(doc.add_paragraph(), f"2. 投資費用：約 {invest_amount:,.0f} 萬元(僅為主機費用，實際金額仍需經廠商報價)。").paragraph_format.first_line_indent = Pt(24)
add_run_kai(doc.add_paragraph(), f"3. 回收年限：{invest_amount:,.0f} 萬元 ÷ {save_money:.1f} 萬元/年 ≒ {payback_year:.1f} 年。").paragraph_format.first_line_indent = Pt(24)

# --- 6. 報告輸出中心 ---
st.markdown("---")
st.subheader("🚀 報告輸出中心")
buf = io.BytesIO(); doc.save(buf); current_word_data = buf.getvalue()
c_b1, c_b2 = st.columns(2)
with c_b1:
    if st.button("🔄 確認數值並同步至打包中心", use_container_width=True):
        if 'report_warehouse' not in st.session_state: st.session_state['report_warehouse'] = {}
        st.session_state['report_warehouse']["4. 冰水主機效益分析"] = current_word_data
        st.success("✅ 數據已鎖定！")
        st.rerun()
with c_b2:
    st.download_button("💾 下載目前的 Word 報告", current_word_data, "冰水主機汰換效益分析.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
