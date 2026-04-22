import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import io

# --- 1. 內部工具：Word 字體格式鎖定 (確保填入的字不會變醜) ---
def fix_run_font(run, size=12):
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)

def safe_replace_tags(doc, data_map):
    """精準替換 Word 中的 {{標籤}} 並保持格式"""
    for p in doc.paragraphs:
        for key, val in data_map.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))
                for run in p.runs:
                    fix_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in data_map.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(val))
                            for run in p.runs:
                                fix_run_font(run)

# --- 2. 主程式開始 ---
st.title("💡 6. 日光燈更換 LED 節能分析")

# 從主程式的 Session State 抓取檔案
uploaded_file = st.session_state.get('global_excel')

if uploaded_file is None:
    st.warning("⚠️ 請先在左側邊欄上傳「完整能源查核 Excel」檔案。")
else:
    # A. 自動讀取與掃描
    all_sheets = pd.read_excel(uploaded_file, sheet_name=None)
    target_lights_list = []

    # 遍歷所有分頁搜尋日光燈
    for sheet_name, df in all_sheets.items():
        if '表九' in sheet_name and '種類' in df.columns:
            # 過濾邏輯：包含 日光燈/螢光燈/T8/T5 且 不含 LED
            mask = (df['種類'].str.contains('日光燈|螢光燈|T8|T5|傳統', na=False)) & \
                   (~df['種類'].str.contains('LED', na=False, case=False))
            found = df[mask].copy()
            if not found.empty:
                found['來源建築物'] = sheet_name
                target_lights_list.append(found)

    if not target_lights_list:
        st.success("✅ 系統掃描完畢：該單位似乎已全面採用 LED，或 Excel 中未發現傳統日光燈資料。")
    else:
        # B. 數據統整與展示
        full_df = pd.concat(target_lights_list, ignore_index=True)
        
        st.info(f"🔍 系統在各分頁中自動偵測到 **{len(full_df)}** 筆待改善燈具資料。")
        
        # 顯示編輯表格 (讓使用者可以微調)
        with st.expander("📝 檢視並微調待改善燈具清單", expanded=True):
            edited_df = st.data_editor(
                full_df[['來源建築物', '種類', '數量(具)', '瓦數(W/具)']],
                use_container_width=True,
                num_rows="dynamic"
            )

        # C. 參數設定區
        st.divider()
        c1, c2, c3 = st.columns(3)
        with c1:
            work_hours = st.number_input("年照明運轉時數 (hr/年)", value=2500)
        with c2:
            # 優先使用主程式算出的平均電費，若無則用預設值
            default_price = st.session_state.get('auto_avg_price', 3.5)
            elec_price = st.number_input("平均電費 (元/度)", value=float(default_price))
        with c3:
            invest_per_unit = st.number_input("每具 LED 更換成本 (元)", value=800)

        # D. 自動效益計算
        # 假設換成 LED 後，瓦數平均降為原本的 45% (例如 T8 92W 換成 40W)
        total_old_kw = (edited_df['瓦數(W/具)'] * edited_df['數量(具)']).sum() / 1000
        total_new_kw = total_old_kw * 0.45 
        saved_kw = total_old_kw - total_new_kw
        saved_kwh = saved_kw * work_hours
        saved_money_year = (saved_kwh * elec_price) / 10000 # 萬元
        
        total_invest = (edited_df['數量(具)'].sum() * invest_per_unit) / 10000 # 萬元
        payback_year = total_invest / saved_money_year if saved_money_year > 0 else 0

        # E. 儀表板顯示
        st.markdown("### 📈 預估節能成效")
        mc1, mc2, mc3, mc4 = st.columns(4)
        mc1.metric("總更換數量", f"{edited_df['數量(具)'].sum():,.0f} 具")
        mc2.metric("年省電量", f"{saved_kwh:,.0f} kWh")
        mc3.metric("年省電費", f"{saved_money_year:.2f} 萬元")
        mc4.metric("回收年限", f"{payback_year:.1f} 年")

        # F. 生成 Word 報告
        st.divider()
        if st.button("📝 生成 LED 汰換專業報告 (docx)", use_container_width=True):
            try:
                # 1. 讀取模板
                # 請確保模板檔案 "有承-10803使用LED燈具及光源.docx" 與 app.py 在同目錄
                doc = Document("有承-10803使用LED燈具及光源.docx")
                
                # 2. 建立標籤映射表 (對應你 docx 裡的標籤)
                # 你可以根據 docx 內容繼續增加 {{標籤}}
                data_map = {
                    "{{NON_LED_COUNT}}": f"{edited_df['數量(具)'].sum():,.0f}",
                    "{{OLD_KW}}": f"{total_old_kw:.2f}",
                    "{{SAVING_KWH}}": f"{saved_kwh:,.0f}",
                    "{{SAVING_MONEY}}": f"{saved_money_year:.2f}",
                    "{{INVEST}}": f"{total_invest:.2f}",
                    "{{PAYBACK}}": f"{payback_year:.1f}",
                    "{{HOURS}}": str(work_hours),
                    "{{E_PRICE}}": f"{elec_price:.2f}"
                }
                
                # 3. 執行替換
                safe_replace_tags(doc, data_map)
                
                # 4. 存入記憶體並提供下載
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # 同時存入倉庫 (供主程式 ZIP 打包用)
                report_name = f"P6_LED更換建議_{edited_df['來源建築物'].iloc[0] if not edited_df.empty else '報告'}"
                st.session_state['report_warehouse'][report_name] = doc_io.getvalue()
                
                st.success(f"✅ 報告已產出並加入下載倉庫！(檔名：{report_name})")
                st.download_button(
                    "📥 下載此份 LED 報告",
                    data=doc_io.getvalue(),
                    file_name=f"{report_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"報告產出失敗，錯誤訊息：{e}")
