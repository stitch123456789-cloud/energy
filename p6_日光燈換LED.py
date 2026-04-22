import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import io
import re

# --- 1. Word 工具函數 ---
def set_font_kai(run, size=12, is_bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = '標楷體'
    run.font.size = Pt(size)
    run.font.bold = is_bold
    run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

# --- 2. 核心讀取與介面 ---
st.title("💡 6. 照明系統 (LED汰換) 自動化報告")

# 從全域抓取已上傳的 Excel 檔案
uploaded_file = st.session_state.get('global_excel')

if uploaded_file is None:
    st.warning("⚠️ 請先在左側邊欄上傳「完整能源查核 Excel」檔案。")
else:
    # 側邊欄參數設定
    st.sidebar.header("⚙️ 照明參數設定")
    work_hours = st.sidebar.number_input("年照明運轉時數 (hr/年)", value=2500, step=100)
    invest_per_unit = st.sidebar.number_input("每具 LED 換裝預估單價 (元)", value=800, step=50)
    
    # 抓取主程式自動算好的電費
    val_from_app = st.session_state.get('auto_avg_price', 3.50)
    electricity_price = st.sidebar.number_input(
        "平均電費 (元/度)", 
        value=float(val_from_app), 
        step=0.1, 
        key="p6_elec"
    )

    # A. 精準提取「表九之四」
    all_sheets = pd.read_excel(uploaded_file, sheet_name=None)
    target_lights_list = []

    for sheet_name, df_raw in all_sheets.items():
        # ✨ 放棄無差別掃描，直接精準鎖定「表九之二」
        clean_name = re.sub(r'\s+', '', sheet_name)
        if bool(re.search(r'表九之二|表9之2|表9-2', clean_name)):
            
            # 定位表頭 (找尋 "燈具種類")
            header_idx = -1
            for idx, row in df_raw.head(10).iterrows():
                row_str = "".join([str(val) for val in row.values if pd.notna(val)])
                if '燈具種類' in row_str and '瓦數' in row_str:
                    header_idx = idx
                    break
            
            if header_idx != -1:
                df = df_raw.iloc[header_idx+1:].copy()
                df.columns = df_raw.iloc[header_idx]
                
                # 確保欄位存在
                if '燈具種類' in df.columns and '瓦數/容量(W/具)' in df.columns and '數量(具)' in df.columns:
                    # 轉數字格式
                    df['數量(具)'] = pd.to_numeric(df['數量(具)'], errors='coerce').fillna(0)
                    df['瓦數/容量(W/具)'] = pd.to_numeric(df['瓦數/容量(W/具)'], errors='coerce').fillna(0)
                    
                    # 過濾：只抓非 LED 的傳統燈具
                    mask = (~df['燈具種類'].str.contains('LED|led', na=False, case=False)) & (df['數量(具)'] > 0)
                    found = df[mask].copy()
                    
                    if not found.empty:
                        found['來源建築物'] = sheet_name
                        target_lights_list.append(found)

    # B. 數據處理與展示
    if not target_lights_list:
        st.success("✅ 該單位「表九之二」已全面採用 LED，或無傳統燈具需汰換。")
    else:
        full_df = pd.concat(target_lights_list, ignore_index=True)
        
        # 計算現況總耗能
        full_df['現況總耗功(W)'] = full_df['瓦數/容量(W/具)'] * full_df['數量(具)']
        total_old_kw = full_df['現況總耗功(W)'].sum() / 1000
        total_count = full_df['數量(具)'].sum()
        
        # 計算改善後效益 (假設換 LED 瓦數剩原本 45%)
        total_new_kw = total_old_kw * 0.45
        saved_kw = total_old_kw - total_new_kw
        saved_kwh = saved_kw * work_hours
        saved_money_wan = (saved_kwh * electricity_price) / 10000
        total_invest_wan = (total_count * invest_per_unit) / 10000
        payback_year = total_invest_wan / saved_money_wan if saved_money_wan > 0 else 0

        st.info(f"🔍 已從「表九之二」精準抓取 **{len(full_df)}** 筆待改善燈具，共計 **{total_count:,.0f}** 具。")
        
        st.write("📋 **待改善傳統燈具清單：**")
        # 顯示給使用者看的精簡表格
        show_df = full_df[['來源建築物', '系統名稱', '燈具種類', '燈管(泡)種類', '瓦數/容量(W/具)', '數量(具)']]
        edited_df = st.data_editor(show_df, use_container_width=True)

        st.markdown("### 📈 預期節能成效預覽")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("預估省電量", f"{saved_kwh:,.0f} 度/年")
        c2.metric("預估省電費", f"{saved_money_wan:.1f} 萬元/年")
        c3.metric("預估投資金額", f"{total_invest_wan:.1f} 萬元")
        c4.metric("回收年限", f"{payback_year:.1f} 年")

        # C. 產出 Word 報告邏輯
        st.divider()
        if st.button("🚀 確認結果並產生 LED 改善報告", use_container_width=True):
            try:
                doc = Document()
                
                # 標題
                doc.add_heading('照明系統改善建議報告', 1)
                
                # 一、 現況說明
                h1 = doc.add_paragraph()
                set_font_kai(h1.add_run('一、現況說明'), size=14, is_bold=True)
                p1 = doc.add_paragraph()
                set_font_kai(p1.add_run("依據申報資料表九之二，貴單位目前仍有部分區域使用傳統螢光燈、T8 或其他較耗能之燈具，共計約 "), size=12)
                set_font_kai(p1.add_run(f"{total_count:,.0f} 具"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p1.add_run("。傳統燈具除了耗電量大之外，亦容易有光衰、閃爍等問題，且維護替換率高。現況待改善清單如下表所示："), size=12)
                
                # 插入現況表格
                table = doc.add_table(rows=1, cols=len(show_df.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(show_df.columns):
                    set_font_kai(hdr_cells[i].paragraphs[0].add_run(col_name), 10, True)
                for _, row in show_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        set_font_kai(row_cells[i].paragraphs[0].add_run(str(val)), 10)
                
                doc.add_paragraph() # 空行
                
                # 二、 改善方案
                h2 = doc.add_paragraph()
                set_font_kai(h2.add_run('二、改善方案'), size=14, is_bold=True)
                p2 = doc.add_paragraph()
                set_font_kai(p2.add_run("建議將上述清單中之傳統燈具，全面汰換為高效率 LED 燈具。LED 燈具具備以下優勢：\n(1) 發光效率高，較傳統燈具節省約 40%~60% 耗能。\n(2) 壽命長達 30,000 小時以上，大幅降低維護與材料成本。\n(3) 無汞、無紫外線，低發熱量，間接降低空調負擔。"), size=12)
                
                # 三、 預期效益
                h3 = doc.add_paragraph()
                set_font_kai(h3.add_run('三、預期效益'), size=14, is_bold=True)
                p3 = doc.add_paragraph()
                set_font_kai(p3.add_run("若將清單中之傳統燈具全數汰換，以年使用時數 "), size=12)
                set_font_kai(p3.add_run(f"{work_hours} 小時"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p3.add_run(" 推估，每年約可節省耗電 "), size=12)
                set_font_kai(p3.add_run(f"{saved_kwh:,.0f} 度/年"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p3.add_run("，節省電費約 "), size=12)
                set_font_kai(p3.add_run(f"{saved_money_wan:.1f} 萬元/年"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p3.add_run("。"), size=12)
                
                p4 = doc.add_paragraph()
                set_font_kai(p4.add_run("預估投資換裝費用約 "), size=12)
                set_font_kai(p4.add_run(f"{total_invest_wan:.1f} 萬元"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p4.add_run("，投資回收年限約為 "), size=12)
                set_font_kai(p4.add_run(f"{payback_year:.1f} 年"), size=12, color=RGBColor(255, 0, 0))
                set_font_kai(p4.add_run(" (實際金額依最終廠商報價為準)。"), size=12)

                # 存檔與打包
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.session_state['report_warehouse']["6. 照明系統改善報告"] = doc_io.getvalue()
                st.success("✅ 報告生成成功！已同步至左側打包中心。")
                
                st.download_button(
                    "💾 單獨下載此份照明報告",
                    data=doc_io.getvalue(),
                    file_name="照明系統汰換效益分析.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"報告生成失敗：{e}")
